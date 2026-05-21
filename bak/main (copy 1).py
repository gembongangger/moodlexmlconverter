import re
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom

def create_moodle_xml(docx_path, output_path):
    # Membaca dokumen
    doc = Document(docx_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # Root elemen Moodle Quiz
    quiz = ET.Element("quiz")

    # Regex untuk memisahkan setiap blok soal
    # Mencari pola [soal no X] sampai soal berikutnya atau akhir dokumen
    question_blocks = re.split(r'\[soal no \d+\]', full_text)
    
    # Ambil judul kategori dari baris pertama (Source 1)
    category_name = doc.paragraphs[0].text if doc.paragraphs else "Imported Questions"

    for block in question_blocks[1:]:  # Skip index 0 karena teks sebelum soal pertama
        lines = [line.strip() for line in block.strip().split('\n') if line.strip()]
        
        if not lines:
            continue

        # Ekstraksi komponen soal
        question_text = ""
        options = []
        keys = []

        for line in lines:
            if line.startswith("isi soal nomor"):
                # Menghapus prefix dan mengambil teks soal
                question_text = re.sub(r'isi soal nomor \d+', '', line).strip()
            elif line.startswith("[opsi"):
                # Mengambil isi opsi setelah tutup kurung ]
                option_content = re.sub(r'\[opsi [A-E]\]', '', line).strip()
                options.append(option_content)
            elif line.startswith("[KUNCI:"):
                # Mencari kunci (B atau B, C, D)
                key_match = re.search(r'\[KUNCI:\s*(.*)\]', line)
                if key_match:
                    keys = [k.strip() for k in key_match.group(1).split(',')]

        # Tentukan tipe soal
        q_type = "multichoice"
        
        # Buat elemen question
        question = ET.SubElement(quiz, "question", type=q_type)
        
        # Nama Soal
        name = ET.SubElement(question, "name")
        ET.SubElement(name, "text").text = f"Soal {len(quiz.findall('question'))}"

        # Teks Soal (mendukung format HTML untuk LaTeX)
        qtext_element = ET.SubElement(question, "questiontext", format="html")
        # Mengubah format matematika sederhana ke format yang dikenal Moodle ( MathJax)
        formatted_qtext = question_text.replace("2x2", "2x^2").replace("cos2x", "\\cos^2 x")
        ET.SubElement(qtext_element, "text").text = f"<![CDATA[<p>{formatted_qtext}</p>]]>"

        # Pengaturan Single/Multiple Answer
        single = "true" if len(keys) <= 1 else "false"
        ET.SubElement(question, "single").text = single
        ET.SubElement(question, "answernumbering").text = "abc"

        # Menambahkan Opsi (Answer)
        map_keys = {'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4}
        num_correct = len(keys)
        
        for i, opt_text in enumerate(options):
            current_label = list(map_keys.keys())[i]
            is_correct = current_label in keys
            
            # Hitung bobot (fraction)
            if is_correct:
                fraction = str(100 / num_correct)
            else:
                fraction = "0" if len(keys) <= 1 else "-50"

            answer = ET.SubElement(question, "answer", fraction=fraction)
            ET.SubElement(answer, "text").text = opt_text
            feedback = ET.SubElement(answer, "feedback")
            ET.SubElement(feedback, "text").text = "Benar" if is_correct else "Salah"

    # Menyimpan file dengan format XML yang rapi
    xml_str = ET.tostring(quiz, encoding='utf-8')
    pretty_xml = minidom.parseString(xml_str).toprettyxml(indent="  ")
    
    # Menghapus escape character CDATA agar terbaca sebagai tag murni di Moodle
    pretty_xml = pretty_xml.replace("&lt;![CDATA[", "<![CDATA[").replace("]]&gt;", "]]>")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)

# Cara penggunaan:
# create_moodle_xml("template.docx", "questions-rev-01.xml")