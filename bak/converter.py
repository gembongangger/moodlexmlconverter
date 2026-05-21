# converter.py
import re
from lxml import etree
import base64
import os
import zipfile
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom


def set_cdata(parent, tag, html):
    el = ET.SubElement(parent, tag)
    el.text = html


def clean_html(html):
    html = re.sub(r'\[begin tabel\]', '', html, flags=re.I)
    html = re.sub(r'\[end tabel\]', '', html, flags=re.I)
    html = re.sub(r'<p>\s*</p>', '', html, flags=re.I)
    return html.strip()


def extract_equations(paragraph):
    """
    OMML (Word Equation) → LaTeX Converter (PRO)
    Support:
    - Pangkat
    - Pecahan
    - Akar
    - Limit
    - Turunan
    - Integral
    - Matriks
    - Vektor
    - Fungsi komposisi
    - Trigonometri
    """
    eqs = []
    xml = paragraph._element.xml.encode("utf-8")
    root = etree.fromstring(xml)

    ns = {
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }

    def parse_node(node):
        if node is None:
            return ""

        tag = etree.QName(node).localname

        # ===== TEKS =====
        if tag == "t":
            return node.text or ""

        # ===== PANGKAT =====
        if tag == "sSup":
            base = node.find(".//m:e", namespaces=ns)
            sup = node.find(".//m:sup", namespaces=ns)
            return f"{parse_node(base)}^{{{parse_node(sup)}}}"

        # ===== PECAHAN =====
        if tag == "f":
            num = node.find(".//m:num", namespaces=ns)
            den = node.find(".//m:den", namespaces=ns)
            return f"\\frac{{{parse_node(num)}}}{{{parse_node(den)}}}"

        # ===== AKAR =====
        if tag == "rad":
            deg = node.find(".//m:deg", namespaces=ns)
            base = node.find(".//m:e", namespaces=ns)
            if deg is not None and parse_node(deg):
                return f"\\sqrt[{parse_node(deg)}]{{{parse_node(base)}}}"
            return f"\\sqrt{{{parse_node(base)}}}"

        # ===== INTEGRAL, SIGMA, N-ARY =====
        if tag == "nary":
            char = node.find(".//m:chr", namespaces=ns)
            sub = node.find(".//m:sub", namespaces=ns)
            sup = node.find(".//m:sup", namespaces=ns)
            expr = node.find(".//m:e", namespaces=ns)

            symbol = parse_node(char)

            if symbol == "∫":
                return f"\\int_{{{parse_node(sub)}}}^{{{parse_node(sup)}}} {parse_node(expr)}"

            if symbol in ["∑", "Σ"]:
                return f"\\sum_{{{parse_node(sub)}}}^{{{parse_node(sup)}}} {parse_node(expr)}"

            return f"{symbol}_{{{parse_node(sub)}}}^{{{parse_node(sup)}}} {parse_node(expr)}"

        # ===== LIMIT =====
        if tag == "limLow":
            base = node.find(".//m:e", namespaces=ns)
            lim = node.find(".//m:lim", namespaces=ns)
            return f"\\lim_{{{parse_node(lim)}}} {parse_node(base)}"

        # ===== TURUNAN (d/dx) =====
        if tag == "d":
            num = node.find(".//m:num", namespaces=ns)
            den = node.find(".//m:den", namespaces=ns)
            return f"\\frac{{{parse_node(num)}}}{{{parse_node(den)}}}"

        # ===== MATRKS =====
        if tag == "m":
            rows = node.findall(".//m:mr", namespaces=ns)
            latex = "\\begin{matrix}"
            for r in rows:
                cells = r.findall(".//m:e", namespaces=ns)
                row_vals = [parse_node(c) for c in cells]
                latex += " & ".join(row_vals) + " \\\\ "
            latex += "\\end{matrix}"
            return latex

        # ===== VEKTOR (panah di atas huruf) =====
        if tag == "acc":
            base = node.find(".//m:e", namespaces=ns)
            return f"\\vec{{{parse_node(base)}}}"

        # ===== FUNGSI KOMPOSISI =====
        if tag == "oMathPara":
            content = "".join(parse_node(c) for c in node)
            content = content.replace("∘", "\\circ ")
            return content

        # ===== GENERIC =====
        result = ""
        for child in node:
            result += parse_node(child)
        return result

    for math in root.findall(".//m:oMath", namespaces=ns):
        latex = parse_node(math)

        # Trigonometri
        latex = latex.replace("sin", "\\sin ")
        latex = latex.replace("cos", "\\cos ")
        latex = latex.replace("tan", "\\tan ")

        # Panah
        latex = latex.replace("→", "\\to ")

        eqs.append(f"\\({latex.strip()}\\)")

    return eqs


def get_docx_images(docx_path):
    images = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            for file in z.namelist():
                if file.startswith('word/media/'):
                    img_name = os.path.basename(file)
                    images[img_name] = z.read(file)
    except Exception as e:
        print(f"Peringatan: Gagal mengekstrak gambar: {e}")
    return images


def table_to_html(table):
    html = '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">'
    for row in table.rows:
        html += '<tr>'
        for cell in row.cells:
            content = cell.text.strip().replace('\n', '<br>')
            html += f'<td style="padding: 8px; border: 1px solid #ccc;">{content}</td>'
        html += '</tr>'
    html += '</table>'
    return html


def create_moodle_xml(docx_path, output_path):
    doc = Document(docx_path)
    images_binary = get_docx_images(docx_path)
    quiz = ET.Element("quiz")

    # Kategori
    category_name = doc.paragraphs[0].text.strip() if doc.paragraphs else "Bank Soal Import"
    cat_q = ET.SubElement(quiz, "question", type="category")
    cat_node = ET.SubElement(cat_q, "category")
    ET.SubElement(cat_node, "text").text = f"$course$/{category_name}"

    questions = []
    current_soal = None

    used_img_keys = sorted(images_binary.keys())
    img_idx = 0

    # Parsing DOCX
    for block in doc.element.body:
        if block.tag.endswith('p'):
            para = [p for p in doc.paragraphs if p._element == block][0]
            text = para.text.strip()

            # Header soal
            match_soal = re.search(r'\[soal no (\d+)\]', text.lower())
            if match_soal:
                if current_soal:
                    questions.append(current_soal)
                current_soal = {
                    'name': match_soal.group(1),
                    'html': '',
                    'options': [],
                    'keys': [],
                    'images': []
                }
                continue

            if current_soal:
                # Opsi
                if re.match(r'(?i)^\[opsi\s+[A-E]\]', text):
                    opt_content = re.sub(r'(?i)^\[opsi\s+[A-E]\]', '', text).strip()
                    current_soal['options'].append(opt_content)

                # Kunci
                elif re.match(r'(?i)^\[kunci\s*:', text):
                    key_raw = re.search(r'(?i)\[kunci\s*:\s*(.*?)\]', text)
                    if key_raw:
                        current_soal['keys'] = [
                            k.strip().upper()
                            for k in key_raw.group(1).split(',')
                        ]

                # Isi soal
                else:
                    eqs = extract_equations(para)
                    html_line = text
                    for eq in eqs:
                        html_line += " " + eq

                    current_soal['html'] += f"<p>{html_line}</p>"

                    # Gambar
                    if 'pic:pic' in para._element.xml or 'w:drawing' in para._element.xml:
                        if img_idx < len(used_img_keys):
                            img_name = used_img_keys[img_idx]
                            current_soal['images'].append({
                                'name': img_name,
                                'data': images_binary[img_name]
                            })
                            current_soal['html'] += (
                                f'<p><img src="@@PLUGINFILE@@/{img_name}" alt="image" /></p>'
                            )
                            img_idx += 1

        # Tabel
        elif block.tag.endswith('tbl'):
            if current_soal:
                table_obj = [t for t in doc.tables if t._element == block][0]
                current_soal['html'] += table_to_html(table_obj)

    if current_soal:
        questions.append(current_soal)

    # Build XML
    for item in questions:
        q = ET.SubElement(quiz, "question", type="multichoice")

        ET.SubElement(ET.SubElement(q, "name"), "text").text = f"Soal {item['name']}"
        qtext_node = ET.SubElement(q, "questiontext", format="html")

        final_html = clean_html(item['html'])
        final_html = final_html.replace("2x2", "2x^2").replace("cos2x", "\\cos^2 x")

        set_cdata(qtext_node, "text", final_html)

        # Gambar ke XML
        for img in item['images']:
            f_tag = ET.SubElement(
                qtext_node, "file",
                name=img['name'], path="/", encoding="base64"
            )
            f_tag.text = base64.b64encode(img['data']).decode('utf-8')

        # Setting jawaban
        is_single = len(item['keys']) <= 1
        ET.SubElement(q, "single").text = "true" if is_single else "false"
        ET.SubElement(q, "answernumbering").text = "abc"
        ET.SubElement(q, "shuffleanswers").text = "1"

        for i, opt_text in enumerate(item['options']):
            label = chr(65 + i)
            correct = label in item['keys']

            if correct:
                fraction = str(100 / len(item['keys']))
            else:
                fraction = "0" if is_single else "-50"

            ans_node = ET.SubElement(q, "answer", fraction=fraction, format="html")
            ET.SubElement(ans_node, "text").text = f"<p>{opt_text}</p>"
            ET.SubElement(ET.SubElement(ans_node, "feedback"), "text").text = (
                "Benar" if correct else "Salah"
            )

    # Simpan file
    xml_data = ET.tostring(quiz, encoding='utf-8')
    pretty_xml = minidom.parseString(xml_data).toprettyxml(indent="  ")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)
