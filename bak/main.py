import re
from lxml import etree
import base64
import os
import zipfile
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom



def set_cdata(parent, tag, html):
    el = ET.SubElement(parent, tag)
    el.text = html

def clean_html(html):
    # Hapus tag tabel custom
    html = re.sub(r'\[begin tabel\]', '', html, flags=re.I)
    html = re.sub(r'\[end tabel\]', '', html, flags=re.I)

    # Hapus paragraf kosong
    html = re.sub(r'<p>\s*</p>', '', html, flags=re.I)

    return html.strip()

def extract_equations(paragraph):
    """
    Ambil equation Word (OMML) dan ubah jadi placeholder LaTeX sederhana
    """
    eqs = []
    xml = paragraph._element.xml.encode("utf-8")
    root = etree.fromstring(xml)

    ns = {
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }

    for math in root.findall('.//m:oMath', namespaces=ns):
        raw = etree.tostring(math, encoding="unicode", method="text")
        latex = raw.replace(' ', '')
        eqs.append(f"\\({latex}\\)")

    return eqs

def get_docx_images(docx_path):
    """Mengekstrak file gambar mentah dari arsip docx."""
    images = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            for file in z.namelist():
                if file.startswith('word/media/'):
                    img_name = os.path.basename(file)
                    images[img_name] = z.read(file)
    except Exception as e:
        print(f"Peringatan: Gagal mengekstrak gambar: {e}")
    return images

def table_to_html(table):
    """Mengonversi tabel Word menjadi format HTML yang bersih untuk Moodle."""
    html = '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">'
    for row in table.rows:
        html += '<tr>'
        for cell in row.cells:
            # Membersihkan teks dalam sel
            content = cell.text.strip().replace('\n', '<br>')
            html += f'<td style="padding: 8px; border: 1px solid #ccc;">{content}</td>'
        html += '</tr>'
    html += '</table>'
    return html

def create_moodle_xml(docx_path, output_path):
    doc = Document(docx_path)
    images_binary = get_docx_images(docx_path)
    quiz = ET.Element("quiz")

    # 1. Tambahkan Kategori (diambil dari baris pertama dokumen)
    category_name = doc.paragraphs[0].text.strip() if doc.paragraphs else "Bank Soal Import"
    cat_q = ET.SubElement(quiz, "question", type="category")
    cat_node = ET.SubElement(cat_q, "category")
    ET.SubElement(cat_node, "text").text = f"$course$/{category_name}"

    questions = []
    current_soal = None
    
    # Melacak index gambar yang sudah digunakan agar tidak tertukar
    used_img_keys = sorted(images_binary.keys())
    img_idx = 0

    # 2. Iterasi Elemen Dokumen (Paragraf & Tabel)
    for block in doc.element.body:
        # Jika elemen adalah Paragraf
        if block.tag.endswith('p'):
            para = [p for p in doc.paragraphs if p._element == block][0]
            text = para.text.strip()
            
            # Deteksi Header Soal: [soal no X]
            match_soal = re.search(r'\[soal no (\d+)\]', text.lower())
            if match_soal:
                if current_soal: questions.append(current_soal)
                current_soal = {
                    'name': match_soal.group(1),
                    'html': '', 
                    'options': [], 
                    'keys': [], 
                    'images': []
                }
                continue
            
            if current_soal is not None:
                # Deteksi Opsi
                if re.match(r'(?i)^\[opsi\s+[A-E]\]', text):
                    opt_content = re.sub(r'(?i)^\[opsi\s+[A-E]\]', '', text).strip()
                    current_soal['options'].append(opt_content)

                # Deteksi Kunci Jawaban
                elif re.match(r'(?i)^\[kunci\s*:', text):
                    key_raw = re.search(r'(?i)\[kunci\s*:\s*(.*?)\]', text)
                    if key_raw:
                        current_soal['keys'] = [k.strip().upper() for k in key_raw.group(1).split(',')]

                # Selain itu = ISI SOAL
                else:
                    eqs = extract_equations(para)

                    html_line = text
                    for eq in eqs:
                        html_line += " " + eq

                    current_soal['html'] += f"<p>{html_line}</p>"


                    # Cek apakah ada gambar di paragraf ini
                    if 'pic:pic' in para._element.xml or 'w:drawing' in para._element.xml:
                        if img_idx < len(used_img_keys):
                            img_name = used_img_keys[img_idx]
                            current_soal['images'].append({
                                'name': img_name,
                                'data': images_binary[img_name]
                            })
                            current_soal['html'] += f'<p><img src="@@PLUGINFILE@@/{img_name}" alt="image" /></p>'
                            img_idx += 1

        # Jika elemen adalah Tabel
        elif block.tag.endswith('tbl'):
            if current_soal is not None:
                table_obj = [t for t in doc.tables if t._element == block][0]
                current_soal['html'] += table_to_html(table_obj)

    if current_soal: questions.append(current_soal)

    # 3. Membangun XML Moodle
    for item in questions:
        q = ET.SubElement(quiz, "question", type="multichoice")
        
        # Nama & Teks Soal
        ET.SubElement(ET.SubElement(q, "name"), "text").text = f"Soal {item['name']}"
        qtext_node = ET.SubElement(q, "questiontext", format="html")
        
        # Format Matematika Dasar (Regex sederhana)
        final_html = clean_html(item['html']).replace("2x2", "2x^2").replace("cos2x", "\\cos^2 x")
        set_cdata(qtext_node, "text", final_html)

        # Lampirkan File Gambar
        for img in item['images']:
            f_tag = ET.SubElement(qtext_node, "file", name=img['name'], path="/", encoding="base64")
            f_tag.text = base64.b64encode(img['data']).decode('utf-8')

        # Pengaturan Jawaban (Single/Multiple)
        is_single = len(item['keys']) <= 1
        ET.SubElement(q, "single").text = "true" if is_single else "false"
        ET.SubElement(q, "answernumbering").text = "abc"
        ET.SubElement(q, "shuffleanswers").text = "1"

        # List Opsi
        for i, opt_text in enumerate(item['options']):
            label = chr(65 + i) # A, B, C, D, E
            correct = label in item['keys']
            
            # Hitung Bobot Nilai
            if correct:
                fraction = str(100 / len(item['keys']))
            else:
                fraction = "0" if is_single else "-50" # Penalti jika salah di pilihan ganda kompleks
                
            ans_node = ET.SubElement(q, "answer", fraction=fraction, format="html")
            ET.SubElement(ans_node, "text").text = f"<![CDATA[<p>{opt_text}</p>]]>"
            ET.SubElement(ET.SubElement(ans_node, "feedback"), "text").text = "Benar" if correct else "Salah"

    # 4. Simpan ke File
    xml_data = ET.tostring(quiz, encoding='utf-8')
    pretty_xml = minidom.parseString(xml_data).toprettyxml(indent="  ")
   
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)

if __name__ == "__main__":
    # Ganti 'template.docx' dengan nama file Anda
    create_moodle_xml("template.docx", "questions-rev-01.xml")
    print("\n--- KONVERSI SELESAI ---")
    print("File Output: questions-rev-01.xml")
    print("Fitur: Teks, Tabel HTML, Gambar Base64, Kunci Ganda.")