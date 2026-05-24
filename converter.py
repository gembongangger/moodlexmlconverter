import re
import base64
import os
import pypandoc
import xml.etree.ElementTree as ET
from xml.dom import minidom
from docx import Document
import zipfile
import openpyxl

def get_docx_images(docx_path):
    images = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            for file in z.namelist():
                if file.startswith('word/media/'):
                    img_name = os.path.basename(file)
                    images[img_name] = base64.b64encode(z.read(file)).decode('utf-8')
    except Exception:
        pass
    return images

def clean_html(html):
    if not html: return ""
    html = re.sub(r'^\s*<\/(?:p|li|h[1-6]|ul|ol|strong|span|em|i|b)>', '', html, flags=re.IGNORECASE)
    html = re.sub(r'<(?:p|li|h[1-6]|ul|ol|strong|span|em|i|b)[^>]*>\s*$', '', html, flags=re.IGNORECASE)
    return html.strip()

def _replace_img_src(text, images_binary_b64):
    imgs = re.findall(r'<img [^>]*src="([^"]+)"', text)
    for src in imgs:
        base_name = os.path.basename(src)
        text = text.replace(src, f'@@PLUGINFILE__/{base_name}')
    return text

def _html_to_text(html):
    return re.sub(r'<[^>]+>', '', html).strip()

def _find_matching_close(html, start, close_tag):
    open_tag = close_tag.replace('/', '')
    tag_base = open_tag.replace('<', '').split()[0]
    depth = 1
    pos = start
    while pos < len(html) and depth > 0:
        if html[pos:].startswith(open_tag) and html[pos+len(open_tag):pos+len(open_tag)+1] in (' ', '>', '\n', '\r', '\t'):
            depth += 1
            pos += len(open_tag)
        elif html[pos:pos+len(close_tag)] == close_tag:
            depth -= 1
            if depth == 0:
                return pos + len(close_tag)
            pos += len(close_tag)
        else:
            pos += 1
    return len(html)

def _preprocess_ol_to_explicit(html):
    """Convert pandoc <ol> lists to explicit numbering for regex parsing.
    Process <ol type='A'> first (innermost), then <ol type='1'> (outermost).
    """
    # Step 1: Process all <ol type="A"> (options) — innermost
    while True:
        m = re.search(r'<ol\s+type="[Aa]"[^>]*>', html)
        if not m: break
        start = m.end()
        end = _find_matching_close(html, start, '</ol>')
        content = html[start:end-len('</ol>')]
        letter = ord('A')
        processed_parts = []
        li_pos = 0
        while li_pos < len(content):
            li_m = re.match(r'<li>(.*?)</li>', content[li_pos:], re.DOTALL)
            if li_m:
                processed_parts.append(f'{chr(letter)}. {li_m.group(1).strip()}')
                letter += 1
                li_pos += li_m.end()
            else:
                processed_parts.append(content[li_pos])
                li_pos += 1
        replacement = ' '.join(processed_parts)
        html = html[:m.start()] + replacement + html[end:]

    # Step 2: Process all <ol type="1"> (questions)
    while True:
        m = re.search(r'<ol(?:\s+start="(\d+)")?\s+type="1"[^>]*>', html)
        if not m: break
        start_num = int(m.group(1)) if m.group(1) else 1
        start = m.end()
        end = _find_matching_close(html, start, '</ol>')
        content = html[start:end-len('</ol>')]
        num = start_num
        processed_parts = []
        li_pos = 0
        while li_pos < len(content):
            li_m = re.match(r'<li>(.*?)</li>', content[li_pos:], re.DOTALL)
            if li_m:
                li_text = li_m.group(1).strip()
                if not li_text.startswith(f'{num}.'):
                    processed_parts.append(f'{num}. {li_text}')
                else:
                    processed_parts.append(li_text)
                num += 1
                li_pos += li_m.end()
            else:
                processed_parts.append(content[li_pos])
                li_pos += 1
        replacement = '\n'.join(processed_parts)
        html = html[:m.start()] + replacement + html[end:]

    return html

def _detect_tf_table(html_block):
    if '<table' not in html_block: return None
    has_benar = re.search(r'BENAR|SALAH', html_block, flags=re.IGNORECASE)
    if not has_benar: return None
    rows = re.findall(r'<tr>(.*?)</tr>', html_block, flags=re.DOTALL)
    if not rows: return None
    header_row = rows[0]
    has_tf_header = ('BENAR' in header_row.upper()) and ('SALAH' in header_row.upper())
    if not has_tf_header or len(rows) < 2: return None
    sub_items = []
    for row in rows[1:]:
        cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row, flags=re.DOTALL)
        if len(cells) >= 2:
            sub_num = _html_to_text(cells[0]).strip()
            sub_text = cells[1].strip()
            sub_items.append({"sub_number": sub_num, "text": sub_text})
    return sub_items if sub_items else None

def _detect_pgk_table(html_block):
    if '<table' not in html_block: return None
    if re.search(r'BENAR|SALAH', html_block, flags=re.IGNORECASE): return None
    if not (re.search(r'Pilihan', html_block, flags=re.IGNORECASE) and re.search(r'Pernyataan', html_block, flags=re.IGNORECASE)):
        return None
    rows = re.findall(r'<tr>(.*?)</tr>', html_block, flags=re.DOTALL)
    if len(rows) < 2: return None
    sub_items = []
    for row in rows[1:]:
        cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row, flags=re.DOTALL)
        if len(cells) >= 2:
            sub_items.append(cells[1].strip())
    return sub_items if sub_items else None

def _parse_tables(full_html, images_binary_b64, questions_list):
    """Find table blocks that follow <ol> questions and parse PGK/TF tables."""
    pos = 0
    while pos < len(full_html):
        table_start = full_html.find('<table', pos)
        if table_start < 0: break

        end = _find_matching_close(full_html, table_start + len('<table'), '</table>')
        table_html = full_html[table_start:end]

        preceding = full_html[:table_start]

        # Detect True/False → matching
        tf_items = _detect_tf_table(table_html)
        if tf_items:
            parent_num = str(len([q for q in questions_list if q.get('type') == 'matching']) + 37)
            ol_num_m = re.findall(r'<ol\s+start="(\d+)"\s+type="1">', preceding)
            if ol_num_m: parent_num = ol_num_m[-1]

            passage = ""
            ol_tag_match = re.search(rf'<ol\s+start="{re.escape(parent_num)}"\s+type="1">.*', preceding, re.DOTALL)
            if ol_tag_match:
                raw = ol_tag_match.group()
                raw = re.sub(r'</?ol[^>]*>', '', raw)
                raw = raw.strip()
                raw_text = re.sub(r'<[^>]+>', '', raw).strip()
                if raw_text and not any(k in raw_text[:20].lower() for k in ['nomor', 'soal nomor', 'petunjuk', 'berilah']):
                    passage = clean_html(raw)
                    passage = _replace_img_src(passage, images_binary_b64)

            sub_questions = []
            for sub in tf_items:
                sub_text = sub['text']
                sub_text = _replace_img_src(sub_text, images_binary_b64)
                sub_questions.append({
                    "text": sub_text,
                    "answer": ""  # will be filled from keys
                })

            questions_list.append({
                "type": "matching",
                "number": parent_num,
                "text": passage if passage else f"Soal {parent_num}",
                "sub_questions": sub_questions,
                "keys": []
            })

            pos = end
            continue

        # Detect PG Kompleks
        pgk_items = _detect_pgk_table(table_html)
        if pgk_items:
            ol_num_m = re.findall(r'<ol\s+start="(\d+)"\s+type="1">', preceding)
            parent_num = ol_num_m[-1] if ol_num_m else str(len([q for q in questions_list if q.get('type') != 'truefalse']) + 1)

            passage = ""
            ol_tag_match = re.search(rf'<ol\s+start="{re.escape(parent_num)}"\s+type="1">.*', preceding, re.DOTALL)
            if ol_tag_match:
                raw = ol_tag_match.group()
                raw = re.sub(r'</?ol[^>]*>', '', raw)
                raw = raw.strip()
                passage = clean_html(raw)
                passage = _replace_img_src(passage, images_binary_b64)

            options = []
            for txt in pgk_items:
                clean_txt = re.sub(r'</?ol[^>]*>', '', txt)
                clean_txt = _replace_img_src(clean_txt, images_binary_b64)
                options.append(clean_html(clean_txt))

            questions_list.append({
                "type": "multichoice",
                "number": parent_num,
                "text": passage if passage else f"Soal {parent_num}",
                "options": options,
                "keys": []
            })
            pos = end
            continue

        pos = end

def extract_questions_dict(docx_path):
    try:
        full_html = pypandoc.convert_file(docx_path, 'html', format='docx', extra_args=['--mathjax', '--wrap=none'])
    except Exception as e:
        return {"questions": [], "category": f"Error: {e}"}

    images_binary_b64 = get_docx_images(docx_path)

    try:
        doc = Document(docx_path)
        category_name = doc.paragraphs[0].text.strip() if doc.paragraphs else "Bank Soal"
    except:
        category_name = "Bank Soal"

    if '<ol type="1">' in full_html or '<ol start=' in full_html:
        questions = _parse_with_ol(full_html, images_binary_b64, category_name)
    else:
        questions = _parse_with_regex(full_html, images_binary_b64, category_name)

    return {
        "category": category_name,
        "questions": questions,
        "images_binary": images_binary_b64
    }

def _parse_with_ol(full_html, images_binary_b64, category_name):
    raw_html = full_html

    # Step 1: Parse tables (PGK and TF) from the raw HTML first
    table_questions = []
    _parse_tables(raw_html, images_binary_b64, table_questions)

    # Step 2: Pre-process OL lists to explicit numbering
    processed = _preprocess_ol_to_explicit(full_html)

    # Step 3: Use regex parser on the processed HTML
    questions_from_regex = _parse_by_regex_inner(processed, images_binary_b64)

    # Step 4: Merge table-based questions (no opt matches found by regex)
    # Identify which questions from tables are NOT already in the regex results
    regex_nums = set(q["number"] for q in questions_from_regex)
    regex_nums_parent = set()
    for q in questions_from_regex:
        if q.get("type") == "truefalse":
            regex_nums_parent.add(q.get("parent_number", q["number"]))
        else:
            regex_nums_parent.add(q["number"])

    for tq in table_questions:
        if tq["type"] == "matching":
            # Replace any regex entry with same number (matching has richer data)
            idx_to_replace = None
            for j, q in enumerate(questions_from_regex):
                if q.get("number") == tq["number"]:
                    idx_to_replace = j
                    break
            if idx_to_replace is not None:
                questions_from_regex[idx_to_replace] = tq
            else:
                questions_from_regex.append(tq)
        elif tq["number"] not in regex_nums:
            questions_from_regex.append(tq)
        elif tq["number"] not in regex_nums_parent:
            questions_from_regex.append(tq)

    return questions_from_regex

def _parse_by_regex_inner(full_html, images_binary_b64):
    """Regex-based parser for pre-processed HTML with explicit numbering."""
    soal_pattern = r'(?:^|\n)\s*(?:<(?:p|div)[^>]*>)?\s*(?:\[soal\s+no\s+\d+\]|(\d+)[\.\)](?:\s|&nbsp;)*|Soal\s*(\d+)|Nomor\s*(\d+))(?:</(?:p|div)>)?'
    soal_matches = list(re.finditer(soal_pattern, full_html, flags=re.IGNORECASE | re.MULTILINE))

    if not soal_matches:
        soal_pattern = r'(?:^|\n)\s*(?:<(?:p|div)[^>]*>)?\s*(\d+)[\.\)](?:\s|&nbsp;)*'
        soal_matches = list(re.finditer(soal_pattern, full_html, flags=re.IGNORECASE | re.MULTILINE))

    questions = []
    current_passage = None

    def detect_shared_passage(text):
        match = re.search(r'Teks\s+ini\s+digunakan\s+untuk\s+menjawab\s+soal\s+nomor\s+(\d+)\s*[–\-]\s*(\d+)', text, flags=re.IGNORECASE)
        if match: return int(match.group(1)), int(match.group(2))
        return None

    for i in range(len(soal_matches)):
        start = soal_matches[i].end()
        end = soal_matches[i+1].start() if i+1 < len(soal_matches) else len(full_html)
        block = full_html[start:end].strip()

        num_val = None
        for g in soal_matches[i].groups():
            if g is not None:
                num_val = int(g)
                break
        if num_val is None:
            num_val = i + 1

        opt_pattern = r'(?:\[opsi\s+[a-e]\]|(?:^|(?<=[>\s]))([A-Ea-e])[\.\)](?:\s|&nbsp;)*)'
        opt_matches = list(re.finditer(opt_pattern, block, flags=re.IGNORECASE | re.MULTILINE))

        if not opt_matches:
            rng = detect_shared_passage(block)
            if rng: current_passage = {"text": block, "start": rng[0], "end": rng[1]}
            continue

        q_text_raw = block[:opt_matches[0].start()]
        q_text_raw = _replace_img_src(q_text_raw, images_binary_b64)

        final_q_text = q_text_raw
        if current_passage and current_passage["start"] <= num_val <= current_passage["end"]:
            p_text = _replace_img_src(current_passage["text"], images_binary_b64)
            final_q_text = f"<div style='background:#f4f4f4; padding:10px; border:1px solid #ddd; border-radius:5px;'>{p_text}</div><br/>{q_text_raw}"

        key_pattern = r'(?:\[kunci\s*:\s*|Kunci\s*(?:Jawaban)?\s*[:\.]\s*|Jawaban\s*[:\.]\s*)(?:<(?:strong|span|p)[^>]*>)?\s*([a-e,\s]+)'
        key_match = re.search(key_pattern, block, flags=re.IGNORECASE)
        keys = [k.strip().upper() for k in re.sub(r'<[^>]+>', '', key_match.group(1)).split(',')] if key_match else []

        options = []
        for j in range(len(opt_matches)):
            opt_start = opt_matches[j].end()
            # Look for next option label or key marker (not HTML tags)
            if j + 1 < len(opt_matches):
                opt_end = opt_matches[j + 1].start()
            else:
                # Look for key marker after last option
                key_marker = re.search(r'(?:\[kunci\s*:|Kunci\s*(?:Jawaban)?\s*[:\.]|Jawaban\s*[:\.])', block[opt_start:], flags=re.IGNORECASE)
                if key_marker:
                    opt_end = opt_start + key_marker.start()
                else:
                    opt_end = len(block)

            clean_opt = block[opt_start:opt_end].strip()
            clean_opt = re.sub(r'^.*?\[opsi\s+[a-e]\]', '', clean_opt, flags=re.IGNORECASE)
            clean_opt = re.sub(r'</?p[^>]*>', '', clean_opt)
            clean_opt = _replace_img_src(clean_opt, images_binary_b64)
            options.append(clean_html(clean_opt))

        questions.append({
            "type": "multichoice",
            "number": str(num_val),
            "text": clean_html(final_q_text),
            "options": options,
            "keys": keys
        })

    return questions

def _parse_with_regex(full_html, images_binary_b64, category_name):
    return _parse_by_regex_inner(full_html, images_binary_b64)

# ---- XML OUTPUT ----

def _embed_images_in_text(txt, imgs_b64):
    txt = re.sub(r'<table', r'<table border="1" style="border-collapse: collapse; width: 100%;"', txt)
    txt = re.sub(r'<td', r'<td style="border: 1px solid black; padding: 5px;"', txt)
    txt = re.sub(r'<th', r'<th style="border: 1px solid black; padding: 5px;"', txt)
    return txt

def _add_images_to_node(node, txt, imgs_b64):
    for img_name in imgs_b64:
        if f'@@PLUGINFILE__/{img_name}' in txt:
            f_tag = ET.SubElement(node, "file", name=img_name, path="/", encoding="base64")
            f_tag.text = imgs_b64[img_name]

def save_multichoice_xml(item, quiz, imgs_b64):
    q = ET.SubElement(quiz, "question", type="multichoice")
    ET.SubElement(ET.SubElement(q, "name"), "text").text = f"Soal {item['number']}"
    qtext_node = ET.SubElement(q, "questiontext", format="html")

    txt = item["text"]
    _add_images_to_node(qtext_node, txt, imgs_b64)
    txt = _embed_images_in_text(txt, imgs_b64)
    ET.SubElement(qtext_node, "text").text = txt

    num_keys = len(item["keys"])
    is_single = num_keys <= 1
    ET.SubElement(q, "single").text = "true" if is_single else "false"
    ET.SubElement(q, "answernumbering").text = "abc"
    ET.SubElement(q, "shuffleanswers").text = "1"

    ET.SubElement(q, "correctfeedback").text = "Jawaban Anda benar."
    ET.SubElement(q, "partiallycorrectfeedback").text = "Sebagian benar."
    ET.SubElement(q, "incorrectfeedback").text = "Jawaban Anda salah."

    for i, opt_text in enumerate(item["options"]):
        label = chr(65 + i)
        correct = label in item["keys"]
        fraction = str(100 / max(1, num_keys)) if correct else ("0" if is_single else str((-1)*100 / (5-max(1, num_keys))))

        ans_node = ET.SubElement(q, "answer", fraction=fraction, format="html")

        _add_images_to_node(ans_node, opt_text, imgs_b64)

        ET.SubElement(ans_node, "text").text = f"<span>{opt_text}</span>"
        ET.SubElement(ET.SubElement(ans_node, "feedback"), "text").text = "Benar" if correct else "Salah"

def save_matching_xml(item, quiz, imgs_b64):
    q = ET.SubElement(quiz, "question", type="matching")
    ET.SubElement(ET.SubElement(q, "name"), "text").text = f"Soal {item['number']}"
    qtext_node = ET.SubElement(q, "questiontext", format="html")

    txt = item["text"]
    _add_images_to_node(qtext_node, txt, imgs_b64)
    txt = _embed_images_in_text(txt, imgs_b64)
    ET.SubElement(qtext_node, "text").text = txt

    ET.SubElement(q, "shuffleanswers").text = "0"

    for sq in item.get("sub_questions", []):
        sub_q = ET.SubElement(q, "subquestion", format="html")
        _add_images_to_node(sub_q, sq["text"], imgs_b64)
        ET.SubElement(sub_q, "text").text = sq["text"]
        ans = ET.SubElement(sub_q, "answer")
        ET.SubElement(ans, "text").text = sq.get("answer", "Benar")

def save_to_xml(data, output_path):
    quiz = ET.Element("quiz")
    cat_q = ET.SubElement(quiz, "question", type="category")
    cat_node = ET.SubElement(cat_q, "category")
    ET.SubElement(cat_node, "text").text = f"$course$/{data['category']}"

    imgs_b64 = data.get("images_binary", {})

    for item in data["questions"]:
        qtype = item.get("type", "multichoice")
        if qtype == "matching":
            save_matching_xml(item, quiz, imgs_b64)
        else:
            save_multichoice_xml(item, quiz, imgs_b64)

    xml_data = ET.tostring(quiz, encoding='utf-8')
    pretty_xml = minidom.parseString(xml_data).toprettyxml(indent="  ")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)

# ---- EXCEL KEY PARSER ----

def parse_keys_from_excel(excel_path, packet_label="PAKET A"):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    keys = {}
    found_header = False

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=False):
        vals = [(cell.column, cell.value) for cell in row]
        non_null = [(c, v) for c, v in vals if v is not None]
        if not non_null:
            continue

        text_concatenated = ' '.join(str(v).strip().upper() for _, v in non_null)

        if packet_label.upper() in text_concatenated:
            found_header = True
            continue

        if not found_header:
            continue

        skip = ['NO. SOAL', 'PERNYATAAN', 'PG.', 'PG KOMPLEKS', 'BENAR-SALAH', 'CATATAN:', 'PAKET A']
        if any(h in text_concatenated for h in skip):
            continue

        # Detect next packet — stop
        if 'PAKET ' in text_concatenated:
            break

        col_data = {c: v for c, v in non_null}

        # PG Tunggal: columns B-C, D-E, F-G, H-I (2-9)
        for col_no in (2, 4, 6, 8):
            col_key = col_no + 1
            if col_no in col_data and col_key in col_data:
                no_val = col_data[col_no]
                key_val = col_data[col_key]
                if isinstance(no_val, (int, float)):
                    no_str = str(int(no_val))
                    key_str = str(key_val).strip().upper()
                    if key_str in 'ABCDE':
                        keys[no_str] = [key_str]

        # PG Kompleks: column J (10)
        if 10 in col_data:
            soal_val = col_data[10]
            if isinstance(soal_val, (int, float)):
                no_str = str(int(soal_val))
                correct_indices = []
                for idx, col in enumerate([11, 12, 13, 14, 15]):
                    if col in col_data:
                        val = str(col_data[col]).strip().upper()
                        if val == 'B':
                            correct_indices.append(chr(65 + idx))
                if correct_indices:
                    keys[no_str] = correct_indices

        # Benar-Salah: column P (16)
        if 16 in col_data:
            soal_val = col_data[16]
            if isinstance(soal_val, (int, float)):
                parent_no = str(int(soal_val))
                for idx, col in enumerate([17, 18, 19, 20]):
                    if col in col_data:
                        key_str = str(col_data[col]).strip().upper()
                        if key_str in ('B', 'S'):
                            sub_no = f"{parent_no}.{idx + 1}"
                            keys[sub_no] = [key_str]

    return keys

def merge_keys(questions, keys_dict):
    merged = []
    for q in questions:
        q = dict(q) if "key_source" not in q else q
        qnum = q["number"]

        if q.get("type") == "matching":
            # Map per sub-question: keys stored as "37.1", "37.2", etc.
            key_mapped = False
            for sidx, sq in enumerate(q.get("sub_questions", [])):
                sub_key = keys_dict.get(f"{qnum}.{sidx + 1}", [])
                if sub_key:
                    sq["answer"] = "Benar" if sub_key[0].upper() == 'B' else "Salah"
                    key_mapped = True
            if key_mapped:
                q["key_source"] = "file"
            elif "key_source" not in q:
                q["key_source"] = "auto"
        elif qnum in keys_dict:
            q["keys"] = keys_dict[qnum]
            q["key_source"] = "file"
        else:
            if "key_source" not in q:
                q["key_source"] = "auto"

        merged.append(q)
    return merged

def create_moodle_xml(docx_path, output_path):
    data = extract_questions_dict(docx_path)
    save_to_xml(data, output_path)
    return {"total_soal": len(data["questions"]), "category": data["category"], "status": "BERHASIL"}
