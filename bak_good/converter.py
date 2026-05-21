import re
import base64
import os
import zipfile
import pypandoc
import tempfile
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom

def extract_content_with_pandoc(para):
    """
    Mengonversi paragraf (teks + rumus) ke HTML/LaTeX menggunakan file sementara 
    agar Pandoc bisa membaca library math OMML secara utuh.
    """
    # Jika paragraf kosong, abaikan
    if not para.text.strip() and not para._element.xpath('.//m:oMath'):
        return ""

    # Buat dokumen docx sementara untuk memicu deteksi Pandoc
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        temp_path = tmp.name

    try:
        # Buat docx minimal yang berisi paragraf ini saja
        new_doc = Document()
        # Salin elemen XML paragraf asli ke dokumen baru
        new_doc.element.body.append(para._element)
        new_doc.save(temp_path)

        # Konversi docx ke HTML dengan flag --mathjax untuk output \( ... \)
        html_output = pypandoc.convert_file(temp_path, 'html', format='docx', extra_args=['--mathjax'])
        
        # Bersihkan tag pembungkus HTML yang tidak perlu
        clean_html = re.sub(r'</?(html|body|head|meta|p|title|span|div)[^>]*>', '', html_output)
        return clean_html.strip()
    
    except Exception as e:
        print(f"Peringatan: Gagal konversi Pandoc: {e}")
        return para.text
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def get_docx_images(docx_path):
    images = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            for file in z.namelist():
                if file.startswith('word/media/'):
                    img_name = os.path.basename(file)
                    images[img_name] = z.read(file)
    except Exception as e:
        print(f"Peringatan: Gagal mengekstrak gambar: {e}")
    return images

def create_moodle_xml(docx_path, output_path):
    doc = Document(docx_path)
    images_binary = get_docx_images(docx_path)
    quiz = ET.Element("quiz")

    # Ambil kategori dari paragraf pertama
    category_name = doc.paragraphs[0].text.strip() if doc.paragraphs else "Bank Soal"
    cat_q = ET.SubElement(quiz, "question", type="category")
    cat_node = ET.SubElement(cat_q, "category")
    ET.SubElement(cat_node, "text").text = f"$course$/{category_name}"

    questions = []
    current_soal = None
    used_img_keys = sorted(images_binary.keys())
    img_idx = 0

    for block in doc.element.body:
        # Cek jika elemen adalah paragraf
        if block.tag.endswith('p'):
            para = [p for p in doc.paragraphs if p._element == block][0]
            text_raw = para.text.strip().lower()

            # Deteksi Header Soal [soal no X]
            match_soal = re.search(r'\[soal no (\d+)\]', text_raw)
            if match_soal:
                if current_soal: questions.append(current_soal)
                current_soal = {'name': match_soal.group(1), 'html': '', 'options': [], 'keys': [], 'images': []}
                continue

            if current_soal:
                # Deteksi Opsi [opsi A]
                if re.match(r'^\[opsi\s+[a-e]\]', text_raw):
                    content = extract_content_with_pandoc(para)
                    # Hapus label [opsi A] dari hasil html
                    content = re.sub(r'(?i)^\[opsi\s+[A-E]\]', '', content).strip()
                    current_soal['options'].append(content)

                # Deteksi Kunci [kunci: A]
                elif re.match(r'^\[kunci\s*:', text_raw):
                    key_raw = re.search(r'(?i)\[kunci\s*:\s*(.*?)\]', para.text)
                    if key_raw:
                        current_soal['keys'] = [k.strip().upper() for k in key_raw.group(1).split(',')]

                # Konten Soal (Teks + Equation)
                else:
                    html_part = extract_content_with_pandoc(para)
                    if html_part:
                        current_soal['html'] += f"<p>{html_part}</p>"

                    # Cek Gambar dalam paragraf
                    if 'pic:pic' in para._element.xml or 'w:drawing' in para._element.xml:
                        if img_idx < len(used_img_keys):
                            img_name = used_img_keys[img_idx]
                            current_soal['images'].append({'name': img_name, 'data': images_binary[img_name]})
                            current_soal['html'] += f'<p><img src="@@PLUGINFILE@@/{img_name}" alt="image" /></p>'
                            img_idx += 1

    if current_soal: questions.append(current_soal)

    # Membangun XML Moodle
    for item in questions:
        q = ET.SubElement(quiz, "question", type="multichoice")
        ET.SubElement(ET.SubElement(q, "name"), "text").text = f"Soal {item['name']}"
        
        qtext_node = ET.SubElement(q, "questiontext", format="html")
        ET.SubElement(qtext_node, "text").text = item['html']

        for img in item['images']:
            f_tag = ET.SubElement(qtext_node, "file", name=img['name'], path="/", encoding="base64")
            f_tag.text = base64.b64encode(img['data']).decode('utf-8')


        # Menghitung jumlah kunci jawaban dan jumlah opsi salah
        num_keys = len(item['keys'])
        num_options = len(item['options'])
        num_wrong = num_options - num_keys

        # Setting jawaban
        is_single = num_keys <= 1
        ET.SubElement(q, "single").text = "true" if is_single else "false"
        ET.SubElement(q, "answernumbering").text = "abc"
        ET.SubElement(q, "shuffleanswers").text = "1"

        for i, opt_html in enumerate(item['options']):
            label = chr(65 + i)
            correct = label in item['keys']
            
            if correct:
                # Logika Skor Benar: 100% dibagi jumlah kunci
                # 1 kunci = 100, 2 kunci = 50, 3 kunci = 33.33333, 4 kunci = 25
                fraction = str(100 / num_keys)
            else:
                # Logika Skor Salah (Penalti):
                if is_single:
                    fraction = "0"
                else:
                    if num_keys == 2:
                        # 2 benar (50% x 2), 3 salah: masing-masing -33.33333%
                        fraction = "-33.33333"
                    elif num_keys == 3:
                        # 3 benar (33.33% x 3), 2 salah: masing-masing -50%
                        fraction = "-50"
                    elif num_keys == 4:
                        # 4 benar (25% x 4), 1 salah: -100%
                        fraction = "-100"
                    else:
                        # Default penalti jika jumlah kunci di luar skenario (misal kunci > 4)
                        fraction = str(-(100 / num_wrong)) if num_wrong > 0 else "0"

            ans_node = ET.SubElement(q, "answer", fraction=fraction, format="html")
            ET.SubElement(ans_node, "text").text = f"<span>{opt_html}</span>"
            ET.SubElement(ET.SubElement(ans_node, "feedback"), "text").text = "Benar" if correct else "Salah"

    # Export ke file
    xml_data = ET.tostring(quiz, encoding='utf-8')
    pretty_xml = minidom.parseString(xml_data).toprettyxml(indent="  ")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)

if __name__ == "__main__":
    # Ganti dengan nama file docx Anda
    create_moodle_xml("template.docx", "hasil_soal.xml")
    print("Konversi selesai! File: hasil_soal.xml")