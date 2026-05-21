import re
import base64
import os
import pypandoc
import xml.etree.ElementTree as ET
from xml.dom import minidom
from docx import Document
import zipfile

def get_docx_images(docx_path):
    """Mengekstrak semua file gambar dari dalam zip docx."""
    images = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            for file in z.namelist():
                if file.startswith('word/media/'):
                    img_name = os.path.basename(file)
                    images[img_name] = z.read(file)
    except Exception as e:
        print(f"Peringatan saat ekstrak gambar: {e}")
    return images

def create_moodle_xml(docx_path, output_path):
    # 1. Konversi seluruh dokumen sekali jalan (Single Heavy Task)
    print("Memulai konversi dokumen dengan Pandoc... (Satu langkah)")
    try:
        # Menggunakan format html dengan mathjax untuk rumus \( ... \)
        full_html = pypandoc.convert_file(docx_path, 'html', format='docx', extra_args=['--mathjax'])
    except Exception as e:
        print(f"Gagal menjalankan Pandoc: {e}")
        return

    # 2. Persiapan Data Gambar
    images_binary = get_docx_images(docx_path)
    
    # 3. Ambil Kategori dari paragraf pertama (menggunakan python-docx)
    doc = Document(docx_path)
    category_name = doc.paragraphs[0].text.strip() if doc.paragraphs else "Bank Soal"

    # 4. Inisialisasi XML Moodle
    quiz = ET.Element("quiz")
    cat_q = ET.SubElement(quiz, "question", type="category")
    cat_node = ET.SubElement(cat_q, "category")
    ET.SubElement(cat_node, "text").text = f"$course$/{category_name}"

    # 5. Parsing Konten menggunakan Regex
    # Split berdasarkan penanda soal
    soal_blocks = re.split(r'(?i)\[soal no \d+\]', full_html)
    soal_numbers = re.findall(r'(?i)\[soal no (\d+)\]', full_html)

    print(f"Ditemukan {len(soal_numbers)} soal. Memproses...")

    for idx, block in enumerate(soal_blocks):
        if idx == 0: continue # Lewati teks sebelum [soal no 1]
        
        num = soal_numbers[idx-1]
        
        # Pisahkan Teks Soal (semua sebelum opsi pertama)
        parts = re.split(r'(?i)\[opsi [a-e]\]', block)
        question_html_raw = parts[0].strip()
        
        # Ambil semua Opsi
        options = re.findall(r'(?s)\[opsi [a-e]\](.*?)(?=\[opsi [a-e]\]|\[kunci\s*:|$)', block, flags=re.IGNORECASE)
        
        # Ambil Kunci
        key_match = re.search(r'(?i)\[kunci\s*:\s*(.*?)\]', block)
        keys = [k.strip().upper() for k in key_match.group(1).split(',')] if key_match else []

        # Membuat Elemen Soal XML
        q = ET.SubElement(quiz, "question", type="multichoice")
        ET.SubElement(ET.SubElement(q, "name"), "text").text = f"Soal {num}"
        
        qtext_node = ET.SubElement(q, "questiontext", format="html")
        
        # --- LOGIKA PEMROSESAN GAMBAR ---
        # Cari tag img yang dihasilkan Pandoc
        img_tags = re.findall(r'<img [^>]*src="([^"]+)"', question_html_raw)
        
        processed_qtext = question_html_raw
        for img_src in img_tags:
            base_name = os.path.basename(img_src)
            
            # Cocokkan dengan binary gambar dari ZIP
            if base_name in images_binary:
                f_tag = ET.SubElement(qtext_node, "file", name=base_name, path="/", encoding="base64")
                f_tag.text = base64.b64encode(images_binary[base_name]).decode('utf-8')
                
                # Update path ke format internal Moodle
                processed_qtext = processed_qtext.replace(img_src, f'@@PLUGINFILE@@/{base_name}')
        
        
	
	# --- TAMBAHKAN LOGIKA BORDER TABEL DISINI ---
        # Mencari tag <table> dan menambahkan atribut border serta style CSS
        processed_qtext = re.sub(
            r'<table', 
            r'<table border="1" style="border-collapse: collapse; width: 100%;"', 
            processed_qtext
        )
        
        # Opsional: Tambahkan border juga pada cell (td) dan header (th)
        processed_qtext = re.sub(r'<td', r'<td style="border: 1px solid black; padding: 5px;"', processed_qtext)
        processed_qtext = re.sub(r'<th', r'<th style="border: 1px solid black; padding: 5px;"', processed_qtext)
        # --- SELESAI ---

        ET.SubElement(qtext_node, "text").text = processed_qtext


        # --- LOGIKA SKOR & OPSI ---
        num_keys = len(keys)
        is_single = num_keys <= 1
        ET.SubElement(q, "single").text = "true" if is_single else "false"
        ET.SubElement(q, "answernumbering").text = "abc"
        ET.SubElement(q, "shuffleanswers").text = "1"

        for i, opt_html in enumerate(options):
            label = chr(65 + i)
            correct = label in keys
            
            # Hitung fraction
            if correct:
                fraction = str(100 / num_keys)
            else:
                # Penalti sederhana jika jawaban jamak
                fraction = "0" if is_single else str((-1)*100 / (5-num_keys))

            ans_node = ET.SubElement(q, "answer", fraction=fraction, format="html")
            ET.SubElement(ans_node, "text").text = f"<span>{opt_html.strip()}</span>"
            ET.SubElement(ET.SubElement(ans_node, "feedback"), "text").text = "Benar" if correct else "Salah"

    # 6. Simpan File
    xml_data = ET.tostring(quiz, encoding='utf-8')
    pretty_xml = minidom.parseString(xml_data).toprettyxml(indent="  ")
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)
    
    print(f"Konversi Berhasil! File tersimpan di: {output_path}")

if __name__ == "__main__":
    # Pastikan file template.docx ada di folder yang sama
    FILE_INPUT = "template.docx" 
    FILE_OUTPUT = "hasil_moodle.xml"
    
    if os.path.exists(FILE_INPUT):
        create_moodle_xml_optimized(FILE_INPUT, FILE_OUTPUT)
    else:
        print(f"Error: File {FILE_INPUT} tidak ditemukan.")